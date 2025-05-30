from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import os
import uuid
import re
from pylatexenc.latex2text import LatexNodes2Text

def generate_word_doc(questions_str, subject, grade):
    # 将LaTeX公式预处理为纯文本
    questions_str = convert_all_latex_to_text(questions_str)
    # 创建临时文件夹
    if not os.path.exists("./tmp"):
        os.makedirs("./tmp")
    
    doc = Document()
    
    # ==================== 设置全局样式 ====================
    # 设置正文字体（中文字体用宋体，西文Times New Roman）
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    normal_style.font.size = Pt(12)  # 小四号字
    normal_style.paragraph_format.line_spacing = 1.5  # 1.5倍行距
    
    # ==================== 设置试卷标题 ====================
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"{grade}{subject}考试试卷")
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title_run.font.size = Pt(22)  # 一号字
    title_run.bold = True
    
    # 添加考试信息行
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run("考试时间：120分钟　　满分：100分")
    info_run.font.name = '楷体'
    info_run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    info_run.font.size = Pt(14)
    info.paragraph_format.space_after = Pt(24)  # 增加标题与正文间距
    
    # ==================== 设置页面布局 ====================
    section = doc.sections[0]
    section.top_margin = Cm(2.5)     # 上边距
    section.bottom_margin = Cm(2.5)  # 下边距
    section.left_margin = Cm(3)      # 左边距
    section.right_margin = Cm(3)     # 右边距
    
    # ==================== 添加试题内容 ====================
    lines = questions_str.strip().split("\n")
    current_question = ""
    current_question_type = None
    
    for line in lines:
        line = line.strip()
        
        # 检测题型标题（一、二、三、等）
        if line.startswith(('一、', '二、', '三、', '四、', '五、')):
            # 处理前一个问题（如果有）
            if current_question:
                add_question_paragraph(doc, current_question, current_question_type)
            
            # 添加题型标题
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)  # 题型前间距
            p.paragraph_format.space_after = Pt(12)   # 题型后间距
            run = p.add_run(line)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(14)
            run.bold = True
            current_question_type = line.split('、')[0]
            current_question = ""
            continue
            
        # 合并连续的问题行（选项和答案）
        if line.startswith(("A.", "B.", "C.", "D.", "1.", "2.", "3.", "4.", "5.","6.","7.","8.","9.","10.","11.","12.","13.",)) or line.endswith(('( )')):
            current_question += "\n" + line
        else:
            # 处理前一个问题（如果有）
            if current_question:
                add_question_paragraph(doc, current_question, current_question_type)
            
            # 开始新问题
            current_question = line
    
    # 处理最后一个问题
    if current_question:
        add_question_paragraph(doc, current_question, current_question_type)
    
    # ==================== 保存文件 ====================
    filename = f"./tmp/{uuid.uuid4()}.docx"
    doc.save(filename)
    return filename

def add_question_paragraph(doc, text, question_type=None):
    """添加问题段落，统一处理格式，并处理LaTeX公式和上下标"""
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    
    # 设置段落格式
    p_format.space_before = Pt(6)  # 段前间距
    p_format.space_after = Pt(6)   # 段后间距
    p_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # 特殊处理选择题选项缩进
    if any(text.startswith(x) for x in ["A.", "B.", "C.", "D."]):
        p_format.left_indent = Cm(0.74)  # 选项缩进
    
    # 处理LaTeX公式
    if contains_latex(text):
        process_latex_content(p, text)
    else:
        # 处理普通文本（包含上下标）
        process_text_with_superscript(p, text)
    
    # 判断题的特殊格式
    if text.endswith(('( )')):
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

def contains_latex(text):
    """检查文本是否包含LaTeX公式"""
    latex_patterns = [
        r'\\[a-zA-Z]+\{',  # 匹配 \command{
        r'\$[^$]+\$',       # 匹配行内公式 $...$
        r'\$\$[^$]+\$\$',   # 匹配块公式 $$...$$
        r'\\\(.*?\\\)',     # 匹配 \(...\)
        r'\\\[.*?\\\]'      # 匹配 \[...\]
    ]
    
    for pattern in latex_patterns:
        if re.search(pattern, text):
            return True
    return False

def process_latex_content(paragraph, text):
    """处理包含LaTeX公式的文本"""
    # 分割文本，分离普通文本和LaTeX公式
    parts = split_text_and_latex(text)
    
    for part in parts:
        if part['type'] == 'text':
            # 处理普通文本（包含上下标）
            process_text_with_superscript(paragraph, part['content'])
        
        elif part['type'] == 'latex':
            # 处理LaTeX公式
            try:
                # 将LaTeX转换为纯文本（保留数学符号）
                converted_text = convert_latex_to_text(part['content'])
                run = paragraph.add_run(converted_text)
                
                # 设置公式样式（斜体，灰色）
                run.font.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)  # 深灰色
                run.font.name = 'Cambria Math'  # Word数学字体
            except Exception as e:
                print(f"公式转换错误: {e}")
                # 转换失败时显示原始LaTeX代码
                run = paragraph.add_run(part['content'])
                run.font.name = 'Courier New'
                run.font.color.rgb = RGBColor(255, 0, 0)  # 红色提示错误

def split_text_and_latex(text):
    """将文本分割为普通文本和LaTeX公式部分"""
    parts = []
    # 匹配常见的LaTeX公式分隔符
    patterns = [
        (r'\$\$(.*?)\$\$', 'block'),    # 块公式 $$...$$
        (r'\$(.*?)\$', 'inline'),       # 行内公式 $...$
        (r'\\\((.*?)\\\)', 'inline'),   # \(...\)
        (r'\\\[(.*?)\\\]', 'block')     # \[...\]
    ]
    
    # 先处理块公式
    for pattern, latex_type in patterns:
        start = 0
        for match in re.finditer(pattern, text, re.DOTALL):
            # 添加前面的普通文本
            if match.start() > start:
                parts.append({
                    'type': 'text',
                    'content': text[start:match.start()]
                })
            
            # 添加公式
            parts.append({
                'type': 'latex',
                'content': match.group(1),
                'latex_type': latex_type
            })
            
            start = match.end()
        
        # 添加剩余文本
        if start < len(text):
            parts.append({
                'type': 'text',
                'content': text[start:]
            })
        
        # 如果找到了公式，返回结果
        if parts:
            return parts
    
    # 如果没有找到公式，返回整个文本作为普通文本
    return [{'type': 'text', 'content': text}]

def convert_latex_to_text(latex_code):
    """将LaTeX公式转换为可显示的文本"""
    # 使用pylatexenc转换基本数学符号
    converter = LatexNodes2Text()
    text = converter.latex_to_text(latex_code)
    
    # 替换特殊数学符号
    replacements = {
        r'\times': '×',
        r'\div': '÷',
        r'\cdot': '·',
        r'\pm': '±',
        r'\mp': '∓',
        r'\leq': '≤',
        r'\geq': '≥',
        r'\neq': '≠',
        r'\approx': '≈',
        r'\infty': '∞',
        r'\alpha': 'α',
        r'\beta': 'β',
        r'\gamma': 'γ',
        r'\delta': 'δ',
        r'\epsilon': 'ε',
        r'\theta': 'θ',
        r'\lambda': 'λ',
        r'\mu': 'μ',
        r'\pi': 'π',
        r'\sigma': 'σ',
        r'\phi': 'φ',
        r'\omega': 'ω',
        r'\sum': '∑',
        r'\prod': '∏',
        r'\int': '∫',
        r'\sqrt': '√',
        r'\frac': '/',
        r'\binom': 'C'
    }
    
    # 应用替换
    for pattern, replacement in replacements.items():
        text = text.replace(pattern, replacement)
    
    # 简化分数表示
    text = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'\1/\2', text)
    
    # 简化开方
    text = re.sub(r'\\sqrt\{([^}]+)\}', r'√\1', text)
    
    return text

def convert_all_latex_to_text(text):
    """将文本中的所有LaTeX公式统一转换为纯文本格式"""
    def replacer(match):
        latex_content = match.group(1)
        return convert_latex_to_text(latex_content)

    # 行内公式 $...$
    text = re.sub(r'\$(.+?)\$', replacer, text)
    # 块公式 $$...$$
    text = re.sub(r'\$\$(.+?)\$\$', replacer, text, flags=re.DOTALL)
    # \( ... \)
    text = re.sub(r'\\\((.+?)\\\)', replacer, text)
    # \[ ... \]
    text = re.sub(r'\\\[(.+?)\\\]', replacer, text, flags=re.DOTALL)

    return text

def process_text_with_superscript(paragraph, text, font_name='Times New Roman', east_asia_font='宋体'):
    """
    处理文本中的上下标格式（如a^2, cm^2, 90^∘）
    支持多种上下标格式：
    1. 常规上下标: a^2, cm^3, 90^∘
    2. 带括号的上下标: a^(2), cm^(3)
    3. 多字符上下标: x^{2y}, a^{12}
    4. 负号上下标: x^{-1}, y^{-2}
    """
    # 替换填空下划线
    if "_____" in text:
        text = text.replace("_____", "________")
    
    # 上下标正则表达式（支持多种格式）
    sup_pattern = r'(\w*[a-zA-Z0-9°])\^(?:(\d+|[+\-]?\w+)|{([^{}]+)}|\(([^()]+)\))'
    
    last_end = 0
    for match in re.finditer(sup_pattern, text):
        # 添加匹配之前的普通文本
        if match.start() > last_end:
            base_text = text[last_end:match.start()]
            run = paragraph.add_run(base_text)
            set_run_font(run, font_name, east_asia_font)
        
        # 提取基体和上标内容
        base = match.group(1)
        sup = match.group(2) or match.group(3) or match.group(4)
        
        # 添加基体
        run_base = paragraph.add_run(base)
        set_run_font(run_base, font_name, east_asia_font)
        
        # 添加上标
        run_sup = paragraph.add_run(sup)
        run_sup.font.superscript = True
        set_run_font(run_sup, font_name, east_asia_font)
        
        last_end = match.end()
    
    # 添加剩余文本
    if last_end < len(text):
        remaining_text = text[last_end:]
        run = paragraph.add_run(remaining_text)
        set_run_font(run, font_name, east_asia_font)

def set_run_font(run, font_name, east_asia_font):
    """设置run的字体格式"""
    run.font.name = font_name
    if east_asia_font:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia_font)